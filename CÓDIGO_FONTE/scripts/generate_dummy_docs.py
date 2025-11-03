from pathlib import Path
from ebooklib import epub
from docx import Document
from fpdf import FPDF

OUT = Path('sample_docs')
OUT.mkdir(exist_ok=True)

for i in range(60):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font('Arial', size=12)
    pdf.cell(0, 10, f'Sample PDF {i}', ln=True)
    pdf.output(str(OUT / f'sample_{i:03d}.pdf'))

for i in range(40):
    book = epub.EpubBook()
    book.set_identifier(str(i))
    book.set_title(f'Sample EPUB {i}')
    book.set_language('pt-BR')
    c1 = epub.EpubHtml(title='Intro', file_name='chap_01.xhtml', lang='pt-BR')
    c1.content = f'<h1>Sample EPUB {i}</h1><p>Conteúdo de exemplo.</p>'
    book.add_item(c1)
    book.toc = (epub.Link('chap_01.xhtml', 'Intro', 'intro'),)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    epub.write_epub(str(OUT / f'sample_epub_{i:03d}.epub'), book)

for i in range(30):
    doc = Document()
    doc.add_heading(f'Sample DOCX {i}', level=1)
    doc.add_paragraph('Parágrafo de exemplo.')
    doc.save(str(OUT / f'sample_{i:03d}.docx'))

for i in range(5):
    (OUT / f'sample_{i:03d}.mobi').write_text('Placeholder MOBI')
for i in range(5,10):
    (OUT / f'sample_{i:03d}.azw').write_text('Placeholder AZW')

print('Arquivos de teste criados em:', OUT)
